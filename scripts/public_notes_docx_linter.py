#!/usr/bin/env python3
"""Lint public lecture-notes DOCX style and student-surface boundaries."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import tempfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Cm, Pt, RGBColor  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

from knowledge_only_rendering_rules import forbidden_advisory_heading_hits, forbidden_advisory_phrase_hits, forbidden_non_knowledge_hits, repeated_template_label_hits

BLUE_RGB = {"0000FF", "0563C1", "2F5496", "1F4E79", "4472C4", "5B9BD5"}
EXPECTED_MARGIN_CM = 2.0
MIN_LINE_SPACING = 1.05
MAX_LINE_SPACING = 1.15
FORBIDDEN_TEXT = {
    "source anchor", "confidence", "evidence score", "recurrence count", "examiner operation", "discriminator axis", "essay plan", "practice question", "answer key", "past paper year", "prediction score", "according to slide", "this slide", "slides say", "source role summary", "source scope", "extraction limitation", "course knowledge map", "predicted essay theme", "section a strategy", "section b strategy", "study order",
}
FORBIDDEN_HEADINGS = {
    "A strong answer should", "Answer Logic", "Conceptual Course Map", "Course Knowledge Map", "Core Exam Claim", "Exam Use", "Common Error / Trap", "Must Master", "Examinable Knowledge Units", "How To Answer", "How To Use This Document", "Knowledge Walkthrough", "Predicted Essay Theme", "Source Role Summary", "Source Scope", "Study Order", "Use This Module",
}


def iter_docx_paths(path: Path) -> list[Path]:
    if path.is_file() and path.suffix.lower() == ".docx":
        return [path]
    if path.is_dir():
        return sorted(child for child in path.rglob("*.docx") if child.is_file())
    return []


def colour_string(colour: Any) -> str | None:
    if colour is None:
        return None
    if getattr(colour, "rgb", None):
        return str(colour.rgb).upper()
    if getattr(colour, "theme_color", None):
        return f"theme:{colour.theme_color}"
    return None


def line_spacing_value(paragraph: Any) -> float | None:
    value = paragraph.paragraph_format.line_spacing
    if value is None and paragraph.style is not None:
        value = paragraph.style.paragraph_format.line_spacing
    if value is None:
        return None
    try:
        return float(value)
    except TypeError:
        return None


def paragraph_has_image(paragraph: Any) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def is_heading(paragraph: Any, visible_index: int) -> bool:
    text = paragraph.text.strip()
    style_name = (paragraph.style.name if paragraph.style else "").casefold()
    if visible_index == 1:
        return True
    if any(token in style_name for token in ["title", "heading", "subheading", "lecture", "module"]):
        return True
    if text.startswith(("Lecture:", "Module:")):
        return True
    return bool(text and len(text) <= 130 and not text.endswith(".") and any(run.bold for run in paragraph.runs))


def lint_docx(path: Path) -> list[dict[str, Any]]:
    failures: list[dict[str, Any]] = []
    doc = Document(path)
    if not doc.sections:
        return [{"type": "missing_docx_section", "path": str(path)}]
    section = doc.sections[0]
    for margin_name in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        value = float(getattr(section, margin_name).cm)
        if abs(value - EXPECTED_MARGIN_CM) > 0.08:
            failures.append({"type": "bad_margin", "path": str(path), "margin": margin_name, "cm": round(value, 3)})
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    lowered = full_text.casefold()
    for phrase in sorted(FORBIDDEN_TEXT):
        if phrase in lowered:
            failures.append({"type": "forbidden_student_text", "path": str(path), "phrase": phrase})
    for heading in sorted(FORBIDDEN_HEADINGS):
        if re.search(rf"(?im)^\s*{re.escape(heading)}\s*:?\s*$", full_text):
            failures.append({"type": "forbidden_internal_heading", "path": str(path), "heading": heading})
    for phrase in forbidden_advisory_phrase_hits(full_text):
        failures.append({"type": "forbidden_advisory_phrase", "path": str(path), "phrase": phrase})
    for heading in forbidden_advisory_heading_hits(full_text):
        failures.append({"type": "forbidden_advisory_heading", "path": str(path), "heading": heading})
    for category in forbidden_non_knowledge_hits(full_text):
        failures.append({"type": "forbidden_non_knowledge_surface", "path": str(path), "category": category})
    for label in repeated_template_label_hits(full_text):
        failures.append({"type": "repeated_rigid_template_label", "path": str(path), "label": label})

    visible_index = 0
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        has_image = paragraph_has_image(paragraph)
        if not text and not has_image:
            continue
        visible_index += 1
        spacing = line_spacing_value(paragraph)
        heading = is_heading(paragraph, visible_index)
        if has_image:
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                failures.append({"type": "image_not_centered", "path": str(path), "paragraph": index})
        elif heading:
            if paragraph.alignment not in {None, WD_ALIGN_PARAGRAPH.LEFT}:
                failures.append({"type": "heading_not_left_aligned", "path": str(path), "paragraph": index})
        else:
            if paragraph.alignment not in {None, WD_ALIGN_PARAGRAPH.LEFT}:
                failures.append({"type": "body_not_left_aligned", "path": str(path), "paragraph": index})
        if spacing is None or (not has_image and not (MIN_LINE_SPACING <= spacing <= MAX_LINE_SPACING)):
            failures.append({"type": "bad_line_spacing", "path": str(path), "paragraph": index, "line_spacing": spacing})
        style_colour = colour_string(paragraph.style.font.color if paragraph.style and paragraph.style.font else None)
        if style_colour and style_colour != "000000":
            failures.append({"type": "style_font_colour_not_black", "path": str(path), "paragraph": index, "colour": style_colour})
        for run_idx, run in enumerate([run for run in paragraph.runs if run.text.strip()], start=1):
            font_names = {name for name in [run.font.name, paragraph.style.font.name if paragraph.style else None] if name}
            if font_names and "Arial" not in font_names:
                failures.append({"type": "non_arial_text", "path": str(path), "paragraph": index, "run": run_idx, "fonts": sorted(font_names)})
            colour = colour_string(run.font.color)
            if colour is not None and colour != "000000":
                failures.append({"type": "run_font_colour_not_black", "path": str(path), "paragraph": index, "run": run_idx, "colour": colour})
            if colour in BLUE_RGB or (colour or "").startswith("theme:"):
                failures.append({"type": "blue_or_theme_text_detected", "path": str(path), "paragraph": index, "run": run_idx, "colour": colour})
    return failures


def lint_path(path: Path) -> dict[str, Any]:
    docx_paths = iter_docx_paths(path)
    failures: list[dict[str, Any]] = []
    if not docx_paths:
        failures.append({"type": "no_docx_found", "path": str(path)})
    for docx_path in docx_paths:
        failures.extend(lint_docx(docx_path))
    return {"pass": not failures, "counts": {"docx_files": len(docx_paths)}, "failures": failures}


def create_bad_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    paragraph = doc.add_paragraph("Course Knowledge Map")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(" This slide says the source role summary should be visible.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 255)
    doc.save(path)


def self_test() -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="public_notes_linter_") as tmp:
        bad = Path(tmp) / "bad.docx"
        create_bad_docx(bad)
        result = lint_path(bad)
        return {"pass": not result["pass"] and bool(result["failures"]), "bad_result": result, "failures": [] if result["failures"] else [{"type": "bad_fixture_not_flagged"}]}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("path", type=Path, nargs="?")
    parser.add_argument("--self-test", action="store_true")
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    try:
        if args.self_test:
            result = self_test()
        elif args.path:
            result = lint_path(args.path)
        else:
            result = {"pass": False, "failures": [{"type": "missing_path_or_self_test"}], "counts": {"docx_files": 0}}
    except Exception as exc:
        result = {"pass": False, "failures": [{"type": "read_error", "error": str(exc)}], "counts": {"docx_files": 0}}
    text = json.dumps(result, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text + "\n", encoding="utf-8")
    else:
        print(text)
    return 0 if result.get("pass") else 1


if __name__ == "__main__":
    raise SystemExit(main())
