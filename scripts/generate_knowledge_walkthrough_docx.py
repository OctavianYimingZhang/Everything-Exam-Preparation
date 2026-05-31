#!/usr/bin/env python3
"""Generate a source-adaptive Knowledge Walkthrough DOCX from a plan."""

from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Cm, Pt, RGBColor  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

from knowledge_only_rendering_rules import (
    forbidden_advisory_heading_hits,
    forbidden_advisory_phrase_hits,
    forbidden_non_knowledge_hits,
    repeated_template_label_hits,
)
from source_scale_budget_rules import floor_for_source_scale

STYLE_DEFAULTS = {
    "route": "knowledge_walkthrough_docx",
    "margin_cm": 2.0,
    "line_spacing": 1.5,
    "body_alignment": "justified",
    "title_alignment": "left",
    "heading_alignment": "left",
    "image_alignment": "center",
    "body_font_pt": 10.5,
    "heading_font_pt": 12.0,
    "lecture_heading_font_pt": 14.0,
    "title_font_pt": 14.0,
    "text_color": "black",
    "module_page_breaks": False,
    "theme_colours_allowed": False,
    "blue_heading_styles_allowed": False,
}

FORBIDDEN_TEXT = {"source anchor", "confidence", "evidence score", "according to slide", "english explanations extracted", "this slide"}


def load_plan(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def style_value(plan: dict[str, Any], key: str) -> Any:
    profile = plan.get("route_docx_style_profile")
    if isinstance(profile, dict) and key in profile:
        return profile[key]
    return STYLE_DEFAULTS[key]


def normalize_run(run, size_pt: float | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size_pt or float(STYLE_DEFAULTS["body_font_pt"]))
    run.font.color.rgb = RGBColor(0, 0, 0)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def paragraph_size(plan: dict[str, Any], kind: str) -> float:
    return {
        "title": float(style_value(plan, "title_font_pt")),
        "lecture": float(style_value(plan, "lecture_heading_font_pt")),
        "heading": float(style_value(plan, "heading_font_pt")),
        "subheading": float(style_value(plan, "heading_font_pt")),
        "body": float(style_value(plan, "body_font_pt")),
    }[kind]


def normalize_paragraph(paragraph, kind: str, plan: dict[str, Any]) -> None:
    paragraph.paragraph_format.line_spacing = float(style_value(plan, "line_spacing"))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if kind == "body" else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        normalize_run(run, paragraph_size(plan, kind))


def set_document_defaults(doc: Document, plan: dict[str, Any]) -> None:
    margin_cm = float(style_value(plan, "margin_cm"))
    section = doc.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    for style_name in ["Normal", "KWTitle", "KWLecture", "KWHeading", "KWSubheading", "KWBody"]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, 1)
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(float(style_value(plan, "body_font_pt")))
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = float(style_value(plan, "line_spacing"))
        style.paragraph_format.space_after = Pt(3)


def add_marked_text(paragraph, text: str, size_pt: float | None = None) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*)", str(text)):
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
        normalize_run(run, size_pt)


def add_paragraph(doc: Document, text: str, plan: dict[str, Any], kind: str = "body"):
    if not str(text or "").strip():
        return None
    style = {"title": "KWTitle", "lecture": "KWLecture", "heading": "KWHeading", "subheading": "KWSubheading", "body": "KWBody"}[kind]
    paragraph = doc.add_paragraph(style=style)
    add_marked_text(paragraph, str(text).strip(), paragraph_size(plan, kind))
    if kind in {"title", "lecture", "heading", "subheading"}:
        for run in paragraph.runs:
            run.bold = True
    normalize_paragraph(paragraph, kind, plan)
    return paragraph


def legacy_to_course_modules(plan: dict[str, Any]) -> list[dict[str, Any]]:
    modules: list[dict[str, Any]] = []
    for lecture in plan.get("legacy_lectures", []) or plan.get("lectures", []) or []:
        if not isinstance(lecture, dict):
            continue
        units = []
        for module in lecture.get("modules", []) or []:
            if not isinstance(module, dict):
                continue
            explanation = " ".join(str(module.get(key) or "").strip() for key in ["module_overview", "knowledge_walkthrough", "key_logic"] if str(module.get(key) or "").strip())
            if explanation:
                units.append({"object_type": "ExaminableKnowledgeUnit", "title": str(module.get("module_title") or "Knowledge unit"), "priority": module.get("priority") or "medium", "explanation": explanation})
        if units:
            modules.append({"object_type": "CourseModule", "module_title": str(lecture.get("lecture_title") or "Course module"), "module_function": str(lecture.get("lecture_overview") or ""), "source_lectures": [str(lecture.get("lecture_id") or lecture.get("lecture_title") or "")], "examinable_units": units})
    return modules


def get_course_modules(plan: dict[str, Any]) -> list[dict[str, Any]]:
    modules = [module for module in plan.get("course_modules", []) or [] if isinstance(module, dict)]
    return modules or legacy_to_course_modules(plan)


def visible_strings(plan: dict[str, Any]) -> list[str]:
    parts = [str(plan.get("title") or ""), str(plan.get("course_knowledge_map") or "")]
    for module in get_course_modules(plan):
        parts.extend(str(module.get(key) or "") for key in ["module_title", "module_function"])
        for unit in module.get("examinable_units", []) or []:
            if isinstance(unit, dict):
                parts.extend(str(unit.get(key) or "") for key in ["title", "explanation", "optional_equation_or_example", "common_confusion_or_boundary"])
    return [part for part in parts if part.strip()]


def star_prefix(priority: Any) -> str:
    mapping = {"high": "★★★", "medium": "★★", "low": "★", "★★★": "★★★", "★★": "★★", "★": "★"}
    value = str(priority or "medium").strip().lower()
    return mapping.get(value, mapping.get(str(priority or "").strip(), "★★"))


def count_public_units(plan: dict[str, Any]) -> int:
    return sum(len([unit for unit in module.get("examinable_units", []) or [] if isinstance(unit, dict)]) for module in get_course_modules(plan))


def validate_source_scale_budget(plan: dict[str, Any]) -> list[str]:
    budget = plan.get("source_scale_budget")
    if not isinstance(budget, dict):
        return ["walkthrough_requires_source_scale_budget"]
    public_units = count_public_units(plan)
    scale_floor = floor_for_source_scale(plan)
    floor_units = int(scale_floor["minimum_public_units"])
    floor_words = int(scale_floor["minimum_visible_words"])
    target_min = budget.get("target_public_units_min")
    if not isinstance(target_min, int) or target_min < 0:
        errors = ["source_scale_budget_missing_target_public_units_min"]
        target_min = floor_units
    elif target_min < floor_units:
        errors = [f"source_scale_budget_target_public_units_min_below_floor:{target_min}<{floor_units}"]
    else:
        errors = []
    target_min = max(int(target_min), floor_units)
    target_words = budget.get("target_words_min")
    if not isinstance(target_words, int) or target_words < 0:
        errors.append("source_scale_budget_missing_target_words_min")
    elif target_words < floor_words:
        errors.append(f"source_scale_budget_target_words_min_below_floor:{target_words}<{floor_words}")
    if budget.get("coverage_floor_status") == "block":
        errors.append("source_scale_budget_blocks_generation")
    if public_units < int(target_min):
        errors.append(f"source_scale_budget_public_units_too_low:{public_units}<{int(target_min)}")
    return errors


def validate_public_density(plan: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    short_units = []
    for module in get_course_modules(plan):
        for unit in module.get("examinable_units", []) or []:
            if not isinstance(unit, dict):
                continue
            explanation = str(unit.get("explanation") or "").strip()
            if len(explanation.split()) < 35:
                short_units.append(str(unit.get("title") or "Knowledge unit"))
    if short_units:
        errors.append("walkthrough_units_too_short_for_reference_caliber:" + "; ".join(short_units[:5]))
    combined = "\n".join(visible_strings(plan))
    word_count = len([token for token in combined.replace("\n", " ").split(" ") if token.strip()])
    budget = plan.get("source_scale_budget")
    target_words_min = budget.get("target_words_min") if isinstance(budget, dict) else None
    scale_floor = floor_for_source_scale(plan)
    effective_target_words = max(target_words_min if isinstance(target_words_min, int) else 0, int(scale_floor["minimum_visible_words"]))
    if effective_target_words > 0 and word_count < effective_target_words:
        errors.append(f"walkthrough_visible_words_too_low:{word_count}<{effective_target_words}")
    return errors


def validate_plan(plan: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if not get_course_modules(plan):
        errors.append("walkthrough_requires_course_modules_or_legacy_lectures")
    errors.extend(validate_source_scale_budget(plan))
    errors.extend(validate_public_density(plan))
    profile = plan.get("route_docx_style_profile")
    if not isinstance(profile, dict):
        errors.append("walkthrough_requires_route_docx_style_profile")
        profile = {}
    if profile.get("route") != "knowledge_walkthrough_docx":
        errors.append("walkthrough_style_profile_wrong_route")
    if profile.get("body_alignment") != "justified":
        errors.append("walkthrough_style_profile_body_not_justified")
    if profile.get("title_alignment") != "left":
        errors.append("walkthrough_style_profile_title_not_left")
    if profile.get("heading_alignment") != "left":
        errors.append("walkthrough_style_profile_heading_not_left")
    if profile.get("image_alignment") != "center":
        errors.append("walkthrough_style_profile_image_not_center")
    if profile.get("text_color") != "black":
        errors.append("walkthrough_style_profile_text_not_black")
    if profile.get("theme_colours_allowed") is not False:
        errors.append("walkthrough_theme_colours_not_disabled")
    if profile.get("blue_heading_styles_allowed") is not False:
        errors.append("walkthrough_blue_heading_styles_not_disabled")
    spacing = profile.get("line_spacing")
    if not isinstance(spacing, (int, float)) or not (1.45 <= float(spacing) <= 1.55):
        errors.append("walkthrough_style_profile_line_spacing_not_1_5")
    for value in visible_strings(plan):
        lowered = value.lower()
        for phrase in FORBIDDEN_TEXT:
            if phrase in lowered:
                errors.append(f"forbidden_text_in_plan:{phrase}")
        for phrase in forbidden_advisory_phrase_hits(value):
            errors.append(f"forbidden_advisory_phrase_in_plan:{phrase}")
        for heading in forbidden_advisory_heading_hits(value):
            errors.append(f"forbidden_advisory_heading_in_plan:{heading}")
        for category in forbidden_non_knowledge_hits(value):
            errors.append(f"forbidden_non_knowledge_surface_in_plan:{category}")
    for label in repeated_template_label_hits("\n".join(visible_strings(plan))):
        errors.append(f"repeated_rigid_template_label:{label}")
    return sorted(set(errors))


def write_docx(plan: dict[str, Any], output_dir: Path, qa_dir: Path, strict: bool) -> dict[str, Any]:
    errors = validate_plan(plan)
    if strict and errors:
        return {"status": "fail", "qa_flags": errors, "documents": []}
    doc = Document()
    set_document_defaults(doc, plan)
    title = str(plan.get("title") or "Lecture Knowledge Walkthrough")
    add_paragraph(doc, title, plan, "title")
    if plan.get("course_knowledge_map"):
        add_paragraph(doc, "Course Knowledge Map", plan, "heading")
        add_paragraph(doc, str(plan["course_knowledge_map"]), plan, "body")
    manifest = {"walkthrough_id": plan.get("walkthrough_id"), "target_group_key": plan.get("target_group_key"), "title": title, "source_scale_budget": plan.get("source_scale_budget"), "public_units": count_public_units(plan), "modules": [], "qa_flags": errors}
    for idx, module in enumerate(get_course_modules(plan)):
        if idx > 0 and bool(style_value(plan, "module_page_breaks")):
            doc.add_page_break()
        module_title = str(module.get("module_title") or "Course module").strip()
        add_paragraph(doc, module_title, plan, "lecture")
        add_paragraph(doc, str(module.get("module_function") or ""), plan, "body")
        core_questions = module.get("core_questions") or []
        if core_questions:
            add_paragraph(doc, "Core questions: " + "; ".join(str(item).strip() for item in core_questions if str(item).strip()), plan, "body")
        module_manifest = {"module_title": module_title, "units": []}
        for unit in module.get("examinable_units", []) or []:
            if not isinstance(unit, dict):
                continue
            unit_title = str(unit.get("title") or "Knowledge unit").strip()
            add_paragraph(doc, f"{star_prefix(unit.get('priority'))} {unit_title}", plan, "subheading")
            add_paragraph(doc, str(unit.get("explanation") or ""), plan, "body")
            add_paragraph(doc, str(unit.get("optional_equation_or_example") or ""), plan, "body")
            add_paragraph(doc, str(unit.get("common_confusion_or_boundary") or ""), plan, "body")
            module_manifest["units"].append({"title": unit_title, "priority": unit.get("priority")})
        manifest["modules"].append(module_manifest)
    output_dir.mkdir(parents=True, exist_ok=True)
    qa_dir.mkdir(parents=True, exist_ok=True)
    filename = "Lecture_Knowledge_Walkthrough.docx"
    docx_path = output_dir / filename
    doc.save(docx_path)
    manifest["documents"] = [{"docx_path": str(docx_path), "filename": filename}]
    manifest["route_docx_style_profile"] = plan.get("route_docx_style_profile") or STYLE_DEFAULTS
    manifest["status"] = "pass" if not errors else "warn"
    (qa_dir / "knowledge_walkthrough_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--plan", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--qa-dir", type=Path)
    parser.add_argument("--clean", action="store_true")
    parser.add_argument("--strict", action="store_true")
    parser.add_argument("--deliverable-only", action="store_true")
    args = parser.parse_args()
    output_dir = args.output_dir
    qa_dir = args.qa_dir or (output_dir if not args.deliverable_only else output_dir.parent / "knowledge_walkthrough_internal_qa")
    if args.clean:
        if output_dir.exists():
            shutil.rmtree(output_dir)
        if qa_dir.exists() and qa_dir != output_dir:
            shutil.rmtree(qa_dir)
    result = write_docx(load_plan(args.plan), output_dir, qa_dir, args.strict)
    print(json.dumps(result, indent=2))
    return 0 if result.get("status") in {"pass", "warn"} else 1


if __name__ == "__main__":
    raise SystemExit(main())
