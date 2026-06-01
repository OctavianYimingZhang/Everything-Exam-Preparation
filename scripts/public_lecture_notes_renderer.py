#!/usr/bin/env python3
"""Render and validate lecture-first public exam-prep notes."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import tempfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Cm, Pt, RGBColor  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

from knowledge_only_rendering_rules import (
    forbidden_advisory_heading_hits,
    forbidden_advisory_phrase_hits,
    forbidden_non_knowledge_hits,
    repeated_template_label_hits,
)
from source_scale_budget_rules import floor_for_source_scale

PUBLIC_STYLE_DEFAULTS = {
    "margin_cm": 2.0,
    "line_spacing": 1.1,
    "body_alignment": "left",
    "title_alignment": "left",
    "heading_alignment": "left",
    "image_alignment": "center",
    "body_font_pt": 10.5,
    "title_font_pt": 14.0,
    "lecture_heading_font_pt": 13.0,
    "module_heading_font_pt": 11.5,
    "block_heading_font_pt": 10.5,
    "text_color": "black",
    "lecture_page_breaks": False,
    "theme_colours_allowed": False,
    "blue_heading_styles_allowed": False,
}

ALLOWED_ROUTES = {"exam_prep_notes_docx", "knowledge_walkthrough_docx"}

KNOWLEDGE_FUNCTIONS = {
    "definition_boundary",
    "mechanism_process",
    "method_readout",
    "graph_data_interpretation",
    "calculation_unit_worked_example",
    "named_example",
    "limitation_trap",
}

BLOCK_TYPES = {
    "definition",
    "mechanism",
    "method",
    "graph_data",
    "calculation",
    "example",
    "limitation",
    "comparison",
    "table",
    "explanation",
}

ROOT_INTERNAL_KEYS = {
    "course_knowledge_map",
    "course_modules",
    "legacy_course_sections",
    "legacy_public_output_points",
    "legacy_lectures",
    "knowledge_cards",
    "exam_overlay_pass",
    "exam_emphasis_profile",
    "question_type_addons",
}

PUBLIC_INTERNAL_KEYS = {
    "source_anchor",
    "source_anchors_visible",
    "source_map",
    "confidence",
    "confidence_band",
    "evidence",
    "evidence_score",
    "recurrence_count",
    "lecture_centrality",
    "examiner_operation",
    "discriminator_axis",
    "exam_specificity",
    "core_exam_claim",
    "exam_use",
    "common_error_or_trap",
    "must_master",
    "practice_question",
    "answer_key",
    "prediction_score",
    "source_role_summary",
    "extraction_limit",
}

FORBIDDEN_PUBLIC_HEADINGS = {
    "A strong answer should",
    "Answer Logic",
    "Course Knowledge Map",
    "Conceptual Course Map",
    "Core Exam Claim",
    "Evidence Used",
    "Exam Specificity",
    "Exam Strategy",
    "Examinable Knowledge Units",
    "Extraction Limitation",
    "How To Answer",
    "How To Answer This Exam",
    "How To Use This Document",
    "Knowledge Sections",
    "Knowledge Walkthrough",
    "Predicted Essay Theme",
    "Recommended Approach",
    "Section A Strategy",
    "Section B Strategy",
    "Source Role Summary",
    "Source Scope",
    "Study Order",
    "Use This Module",
    "What This Lecture Is About",
    "What This Module Explains",
}

FORBIDDEN_PUBLIC_PHRASES = {
    "course knowledge map",
    "conceptual course map",
    "extraction limitation",
    "predicted essay theme",
    "section a strategy",
    "section b strategy",
    "source role summary",
    "source scope",
    "study order",
}

CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
WORD_RE = re.compile(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?")


def load_plan(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def style_value(plan: dict[str, Any], key: str, route: str) -> Any:
    profile = plan.get("route_docx_style_profile")
    if isinstance(profile, dict) and key in profile:
        return profile[key]
    if key == "route":
        return route
    return PUBLIC_STYLE_DEFAULTS[key]


def normalize_run(run: Any, size_pt: float) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def paragraph_size(plan: dict[str, Any], kind: str, route: str) -> float:
    key = {
        "title": "title_font_pt",
        "lecture": "lecture_heading_font_pt",
        "module": "module_heading_font_pt",
        "block_heading": "block_heading_font_pt",
        "body": "body_font_pt",
    }[kind]
    return float(style_value(plan, key, route))


def normalize_paragraph(paragraph: Any, kind: str, plan: dict[str, Any], route: str) -> None:
    paragraph.paragraph_format.line_spacing = float(style_value(plan, "line_spacing", route))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    if kind == "body":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        normalize_run(run, paragraph_size(plan, kind, route))


def ensure_style(doc: Document, name: str, size_pt: float, line_spacing: float) -> None:
    if name not in doc.styles:
        doc.styles.add_style(name, 1)
    style = doc.styles[name]
    style.font.name = "Arial"
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)


def set_document_defaults(doc: Document, plan: dict[str, Any], route: str) -> None:
    margin_cm = float(style_value(plan, "margin_cm", route))
    section = doc.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    line_spacing = float(style_value(plan, "line_spacing", route))
    ensure_style(doc, "EPNTitle", paragraph_size(plan, "title", route), line_spacing)
    ensure_style(doc, "EPNLectureHeading", paragraph_size(plan, "lecture", route), line_spacing)
    ensure_style(doc, "EPNModuleHeading", paragraph_size(plan, "module", route), line_spacing)
    ensure_style(doc, "EPNBlockHeading", paragraph_size(plan, "block_heading", route), line_spacing)
    ensure_style(doc, "EPNBody", paragraph_size(plan, "body", route), line_spacing)
    ensure_style(doc, "Normal", paragraph_size(plan, "body", route), line_spacing)


def add_marked_text(paragraph: Any, text: str, size_pt: float) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*)", str(text)):
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
        normalize_run(run, size_pt)


def add_paragraph(doc: Document, text: str, plan: dict[str, Any], route: str, kind: str = "body") -> Any | None:
    text = str(text or "").strip()
    if not text:
        return None
    style = {
        "title": "EPNTitle",
        "lecture": "EPNLectureHeading",
        "module": "EPNModuleHeading",
        "block_heading": "EPNBlockHeading",
        "body": "EPNBody",
    }[kind]
    paragraph = doc.add_paragraph(style=style)
    add_marked_text(paragraph, text, paragraph_size(plan, kind, route))
    if kind != "body":
        for run in paragraph.runs:
            run.bold = True
    normalize_paragraph(paragraph, kind, plan, route)
    return paragraph


def walk_values(value: Any) -> list[tuple[str, Any]]:
    found: list[tuple[str, Any]] = []
    if isinstance(value, dict):
        for key, child in value.items():
            found.append((str(key), child))
            found.extend(walk_values(child))
    elif isinstance(value, list):
        for child in value:
            found.extend(walk_values(child))
    return found


def visible_text_parts(plan: dict[str, Any]) -> list[str]:
    parts = [str(plan.get("title") or "")]
    for lecture in plan.get("public_lecture_sections", []) or []:
        if not isinstance(lecture, dict):
            continue
        parts.extend(str(lecture.get(key) or "") for key in ["lecture_title", "lecture_scope"])
        for module in lecture.get("modules", []) or []:
            if not isinstance(module, dict):
                continue
            parts.extend(str(module.get(key) or "") for key in ["module_title", "explanation"])
            for block in module.get("blocks", []) or []:
                if isinstance(block, dict):
                    parts.extend(str(block.get(key) or "") for key in ["label", "content"])
                else:
                    parts.append(str(block))
    return [part for part in parts if part.strip()]


def visible_text(plan: dict[str, Any]) -> str:
    return "\n".join(visible_text_parts(plan))


def visible_word_count(plan: dict[str, Any]) -> int:
    return len(WORD_RE.findall(visible_text(plan)))


def count_public_modules(plan: dict[str, Any]) -> int:
    count = 0
    for lecture in plan.get("public_lecture_sections", []) or []:
        if isinstance(lecture, dict):
            count += len([module for module in lecture.get("modules", []) or [] if isinstance(module, dict)])
    return count


def language_allows_cjk(profile: Any) -> bool:
    if not isinstance(profile, dict):
        return False
    language = str(profile.get("output_language") or profile.get("language") or "English").strip().casefold()
    return bool(profile.get("allow_bilingual") or profile.get("allow_non_english") or language in {"chinese", "zh", "zh-cn", "bilingual", "mixed"})


def _validate_heading_text(text: str, errors: list[str], where: str) -> None:
    normalized = text.strip().rstrip(":").casefold()
    for heading in FORBIDDEN_PUBLIC_HEADINGS:
        if normalized == heading.casefold():
            errors.append(f"forbidden_public_heading:{where}:{heading}")


def _validate_block(block: Any, errors: list[str], where: str) -> None:
    if not isinstance(block, dict):
        errors.append(f"public_block_not_object:{where}")
        return
    block_type = str(block.get("block_type") or "").strip()
    content = block.get("content")
    if block_type and block_type not in BLOCK_TYPES:
        errors.append(f"public_block_bad_type:{where}:{block_type}")
    if content is None or (isinstance(content, str) and not content.strip()) or (isinstance(content, list) and not content):
        errors.append(f"public_block_missing_content:{where}")


def _validate_public_module(module: Any, errors: list[str], where: str) -> None:
    if not isinstance(module, dict):
        errors.append(f"public_module_not_object:{where}")
        return
    title = str(module.get("module_title") or "").strip()
    explanation = str(module.get("explanation") or "").strip()
    functions = module.get("knowledge_functions")
    if not title:
        errors.append(f"public_module_missing_title:{where}")
    else:
        _validate_heading_text(title, errors, where)
    if len(explanation) < 160:
        errors.append(f"public_module_explanation_too_short:{where}")
    if not isinstance(functions, list):
        errors.append(f"public_module_missing_knowledge_functions:{where}")
        normalized_functions: set[str] = set()
    else:
        normalized_functions = {str(item).strip() for item in functions if str(item).strip()}
        if len(normalized_functions) < 2:
            errors.append(f"public_module_depth_functions_too_few:{where}")
        bad = sorted(function for function in normalized_functions if function not in KNOWLEDGE_FUNCTIONS)
        for function in bad:
            errors.append(f"public_module_bad_knowledge_function:{where}:{function}")
    blocks = module.get("blocks")
    if not isinstance(blocks, list) or not blocks:
        errors.append(f"public_module_missing_blocks:{where}")
    else:
        for index, block in enumerate(blocks, start=1):
            _validate_block(block, errors, f"{where}.blocks[{index}]")


def _validate_public_lectures(plan: dict[str, Any], errors: list[str]) -> None:
    lectures = plan.get("public_lecture_sections")
    if not isinstance(lectures, list) or not lectures:
        errors.append("requires_public_lecture_sections")
        return
    for lecture_index, lecture in enumerate(lectures, start=1):
        where = f"public_lecture_sections[{lecture_index}]"
        if not isinstance(lecture, dict):
            errors.append(f"public_lecture_not_object:{where}")
            continue
        title = str(lecture.get("lecture_title") or "").strip()
        if not title:
            errors.append(f"public_lecture_missing_title:{where}")
        else:
            _validate_heading_text(title, errors, where)
        modules = lecture.get("modules")
        if not isinstance(modules, list) or not modules:
            errors.append(f"public_lecture_missing_modules:{where}")
            continue
        for module_index, module in enumerate(modules, start=1):
            _validate_public_module(module, errors, f"{where}.modules[{module_index}]")


def _validate_source_scale(plan: dict[str, Any], errors: list[str]) -> None:
    budget = plan.get("source_scale_budget")
    if not isinstance(budget, dict):
        errors.append("requires_source_scale_budget")
        return
    if budget.get("coverage_floor_status") == "block":
        errors.append("source_scale_budget_blocks_generation")
    scale_floor = floor_for_source_scale(plan)
    floor_units = int(scale_floor["minimum_public_units"])
    floor_words = int(scale_floor["minimum_visible_words"])
    target_public_units_min = budget.get("target_public_units_min")
    if isinstance(target_public_units_min, int) and target_public_units_min < floor_units:
        errors.append(f"target_public_units_min_below_source_scale_floor:{target_public_units_min}<{floor_units}")
    if isinstance(target_public_units_min, int) and count_public_modules(plan) < target_public_units_min:
        errors.append(f"source_scale_public_modules_too_low:{count_public_modules(plan)}<{target_public_units_min}")
    target_words_min = budget.get("target_words_min")
    if isinstance(target_words_min, int) and target_words_min < floor_words:
        errors.append(f"target_words_min_below_source_scale_floor:{target_words_min}<{floor_words}")
    if count_public_modules(plan) < floor_units:
        errors.append(f"source_scale_public_modules_below_floor:{count_public_modules(plan)}<{floor_units}")
    if visible_word_count(plan) < floor_words:
        errors.append(f"source_scale_words_below_floor:{visible_word_count(plan)}<{floor_words}")
    if isinstance(target_words_min, int) and visible_word_count(plan) < target_words_min:
        errors.append(f"source_scale_words_too_low:{visible_word_count(plan)}<{target_words_min}")


def _validate_style(plan: dict[str, Any], route: str, errors: list[str]) -> None:
    profile = plan.get("route_docx_style_profile")
    if not isinstance(profile, dict):
        errors.append("requires_route_docx_style_profile")
        profile = {}
    if profile.get("route") != route:
        errors.append(f"style_profile_wrong_route:{profile.get('route')}!={route}")
    if abs(float(profile.get("margin_cm", -1)) - 2.0) > 0.08:
        errors.append("style_profile_margin_not_2_0")
    spacing = profile.get("line_spacing")
    if not isinstance(spacing, (int, float)) or not (1.05 <= float(spacing) <= 1.15):
        errors.append("style_profile_line_spacing_not_compact")
    if profile.get("body_alignment") != "left":
        errors.append("style_profile_body_not_left")
    if profile.get("title_alignment") != "left":
        errors.append("style_profile_title_not_left")
    if profile.get("heading_alignment") != "left":
        errors.append("style_profile_heading_not_left")
    if profile.get("image_alignment") != "center":
        errors.append("style_profile_image_not_center")
    if profile.get("text_color") != "black":
        errors.append("style_profile_text_not_black")
    if profile.get("theme_colours_allowed") is not False:
        errors.append("style_profile_theme_colours_allowed")
    if profile.get("blue_heading_styles_allowed") is not False:
        errors.append("style_profile_blue_headings_allowed")


def validate_public_lecture_notes_plan(plan: dict[str, Any], route: str) -> list[str]:
    errors: list[str] = []
    if route not in ALLOWED_ROUTES:
        errors.append(f"unsupported_route:{route}")
    if plan.get("object_type") != "PublicLectureNotesPlan":
        errors.append("plan_object_type_not_public_lecture_notes_plan")
    for key in ROOT_INTERNAL_KEYS:
        if key in plan:
            errors.append(f"internal_root_key_not_public:{key}")
    _validate_source_scale(plan, errors)
    _validate_style(plan, route, errors)
    _validate_public_lectures(plan, errors)

    language_profile = plan.get("output_language_profile")
    if not isinstance(language_profile, dict):
        errors.append("requires_output_language_profile")
    text = visible_text(plan)
    if not language_allows_cjk(language_profile) and CJK_RE.search(text):
        errors.append("default_english_public_text_contains_cjk")
    lowered = text.casefold()
    for phrase in sorted(FORBIDDEN_PUBLIC_PHRASES):
        if phrase in lowered:
            errors.append(f"forbidden_public_phrase:{phrase}")
    for heading in forbidden_advisory_heading_hits(text):
        errors.append(f"forbidden_advisory_heading:{heading}")
    for phrase in forbidden_advisory_phrase_hits(text):
        errors.append(f"forbidden_advisory_phrase:{phrase}")
    for category in forbidden_non_knowledge_hits(text):
        errors.append(f"forbidden_non_knowledge_surface:{category}")
    for label in repeated_template_label_hits(text):
        errors.append(f"repeated_rigid_template_label:{label}")
    for key, value in walk_values(plan.get("public_lecture_sections")):
        if key in PUBLIC_INTERNAL_KEYS:
            errors.append(f"forbidden_internal_key_in_public_surface:{key}")
        if isinstance(value, str):
            _validate_heading_text(value, errors, f"public_text:{key}")
    return sorted(set(errors))


def _block_text(block: dict[str, Any]) -> list[str]:
    label = str(block.get("label") or "").strip()
    content = block.get("content")
    if isinstance(content, list):
        rendered = "; ".join(str(item).strip() for item in content if str(item).strip())
    else:
        rendered = str(content or "").strip()
    if label:
        return [f"{label}: {rendered}"] if rendered else [label]
    return [rendered] if rendered else []


def write_public_lecture_notes_docx(
    plan: dict[str, Any],
    route: str,
    output_dir: Path,
    qa_dir: Path,
    strict: bool,
    manifest_name: str,
) -> dict[str, Any]:
    errors = validate_public_lecture_notes_plan(plan, route)
    if strict and errors:
        return {"status": "fail", "qa_flags": errors, "documents": []}

    doc = Document()
    set_document_defaults(doc, plan, route)
    title = str(plan.get("title") or "Lecture Knowledge Walkthrough").strip()
    add_paragraph(doc, title, plan, route, "title")
    manifest: dict[str, Any] = {
        "route": route,
        "notes_plan_id": plan.get("notes_plan_id") or plan.get("walkthrough_id"),
        "target_group_key": plan.get("target_group_key"),
        "title": title,
        "source_scale_budget": plan.get("source_scale_budget"),
        "public_modules": count_public_modules(plan),
        "visible_words": visible_word_count(plan),
        "qa_flags": errors,
        "lectures": [],
    }

    for lecture_index, lecture in enumerate(plan.get("public_lecture_sections", []) or [], start=1):
        if not isinstance(lecture, dict):
            continue
        if lecture_index > 1 and bool(style_value(plan, "lecture_page_breaks", route)):
            doc.add_page_break()
        lecture_title = str(lecture.get("lecture_title") or f"Lecture {lecture_index}").strip()
        add_paragraph(doc, lecture_title, plan, route, "lecture")
        add_paragraph(doc, str(lecture.get("lecture_scope") or ""), plan, route, "body")
        lecture_manifest = {"lecture_title": lecture_title, "modules": []}
        for module in lecture.get("modules", []) or []:
            if not isinstance(module, dict):
                continue
            module_title = str(module.get("module_title") or "Knowledge module").strip()
            add_paragraph(doc, module_title, plan, route, "module")
            add_paragraph(doc, str(module.get("explanation") or ""), plan, route, "body")
            for block in module.get("blocks", []) or []:
                if isinstance(block, dict):
                    for text in _block_text(block):
                        add_paragraph(doc, text, plan, route, "body")
            lecture_manifest["modules"].append(
                {
                    "module_title": module_title,
                    "knowledge_functions": module.get("knowledge_functions") or [],
                    "block_count": len(module.get("blocks") or []),
                }
            )
        manifest["lectures"].append(lecture_manifest)

    output_dir.mkdir(parents=True, exist_ok=True)
    qa_dir.mkdir(parents=True, exist_ok=True)
    filename = "Lecture_Knowledge_Walkthrough.docx"
    docx_path = output_dir / filename
    doc.save(docx_path)
    manifest["documents"] = [{"docx_path": str(docx_path), "filename": filename}]
    manifest["route_docx_style_profile"] = plan.get("route_docx_style_profile") or {
        "route": route,
        **PUBLIC_STYLE_DEFAULTS,
    }
    manifest["status"] = "pass" if not errors else "warn"
    (qa_dir / manifest_name).write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def sample_plan(route: str = "exam_prep_notes_docx") -> dict[str, Any]:
    return {
        "object_type": "PublicLectureNotesPlan",
        "notes_plan_id": "sample_public_lecture_notes",
        "title": "Lecture Knowledge Walkthrough",
        "target_group_key": "sample_course",
        "source_scale_budget": {
            "source_units_count": 2,
            "readable_source_blocks": 8,
            "protected_knowledge_units_total": 4,
            "excluded_non_knowledge_units_total": 1,
            "target_public_units_min": 8,
            "target_words_min": 420,
            "compression_mode": "explain_not_dump",
            "coverage_floor_status": "pass",
        },
        "output_language_profile": {"output_language": "English", "allow_bilingual": False},
        "route_docx_style_profile": {
            "route": route,
            "margin_cm": 2.0,
            "line_spacing": 1.1,
            "body_alignment": "left",
            "title_alignment": "left",
            "heading_alignment": "left",
            "image_alignment": "center",
            "body_font_pt": 10.5,
            "title_font_pt": 14.0,
            "lecture_heading_font_pt": 13.0,
            "module_heading_font_pt": 11.5,
            "text_color": "black",
            "lecture_page_breaks": False,
            "theme_colours_allowed": False,
            "blue_heading_styles_allowed": False,
        },
        "public_lecture_sections": [
            {
                "lecture_title": "Lecture 1: Membranes as selective chemical boundaries",
                "modules": [
                    {
                        "module_title": "Bilayer chemistry explains selective permeability",
                        "knowledge_functions": ["definition_boundary", "mechanism_process"],
                        "explanation": "A phospholipid bilayer is a chemical boundary because hydrophilic head groups face water while hydrophobic tails form an internal non-polar core. Small non-polar molecules cross more easily than ions because charged solutes must shed hydration shells and enter an energetically unfavourable core.",
                        "blocks": [
                            {
                                "block_type": "limitation",
                                "content": "Permeability is not absolute: transport proteins change the effective boundary for specific ions or polar solutes.",
                            }
                        ],
                    },
                    {
                        "module_title": "Fluidity controls how membrane proteins can move and interact",
                        "knowledge_functions": ["mechanism_process", "limitation_trap"],
                        "explanation": "Membrane fluidity describes lateral movement within the bilayer and depends on lipid packing. Unsaturated tails reduce packing, while cholesterol buffers extremes by limiting movement at high temperature and preventing tight packing at low temperature.",
                        "blocks": [{"block_type": "limitation", "content": "Fluidity does not mean the membrane is disorganised; it means movement is constrained enough to preserve a boundary while allowing protein interactions."}],
                    },
                    {
                        "module_title": "Selective permeability creates a need for transport proteins",
                        "knowledge_functions": ["definition_boundary", "named_example"],
                        "explanation": "Selective permeability means the bilayer allows some substances through faster than others. Oxygen and carbon dioxide diffuse readily, whereas ions and most polar metabolites need channels, carriers or pumps because the hydrophobic core creates an energetic barrier.",
                        "blocks": [{"block_type": "example", "content": "Ion channels illustrate the boundary because they create a protein-lined route through a membrane that the ion cannot cross efficiently by dissolving in the lipid core."}],
                    },
                    {
                        "module_title": "Compartment membranes let cells separate incompatible chemistry",
                        "knowledge_functions": ["mechanism_process", "named_example"],
                        "explanation": "Internal membranes allow different reactions to occur in different chemical environments. This matters because enzymes, ion concentrations and redox conditions can be tuned locally rather than averaged across the whole cytoplasm.",
                        "blocks": [{"block_type": "example", "content": "An acidic lumen and a neutral cytosol can support different reaction sets because the membrane restricts uncontrolled mixing."}],
                    },
                ],
            },
            {
                "lecture_title": "Lecture 2: Transport proteins convert gradients into flux",
                "modules": [
                    {
                        "module_title": "Electrochemical gradients combine concentration and voltage",
                        "knowledge_functions": ["mechanism_process", "calculation_unit_worked_example"],
                        "explanation": "An electrochemical gradient has a concentration component and an electrical component. A cation may move down its concentration gradient but against the membrane voltage, so the direction of net movement depends on the combined free-energy difference rather than concentration alone.",
                        "blocks": [
                            {
                                "block_type": "calculation",
                                "label": "Worked example",
                                "content": "For an ion, compare the chemical term with the voltage term before predicting net flux direction.",
                            }
                        ],
                    },
                    {
                        "module_title": "Channels and carriers differ in how they limit flux",
                        "knowledge_functions": ["definition_boundary", "method_readout"],
                        "explanation": "Channels form pores that allow rapid movement when open, while carriers bind solute and change conformation. The distinction matters experimentally because channels show high conductance and gating, whereas carriers show saturable transport as binding sites become occupied.",
                        "blocks": [{"block_type": "comparison", "content": "A channel is limited mainly by opening and ion passage; a carrier is limited by binding, conformational change and release."}],
                    },
                    {
                        "module_title": "Active transport couples uphill movement to an energy source",
                        "knowledge_functions": ["mechanism_process", "limitation_trap"],
                        "explanation": "Active transport moves a solute against its electrochemical gradient by coupling transport to ATP hydrolysis, light, redox energy or another gradient. The coupled reaction makes the overall free-energy change favourable even when the solute step alone is unfavourable.",
                        "blocks": [{"block_type": "limitation", "content": "Uphill movement cannot be inferred from concentration alone; the membrane voltage and coupled energy source must also be considered."}],
                    },
                    {
                        "module_title": "Transport measurements need controls for leak and driving force",
                        "knowledge_functions": ["method_readout", "graph_data_interpretation"],
                        "explanation": "A transport assay is interpretable only when the driving force and leak background are defined. Controls without transporter, without substrate or without energy source distinguish transporter-dependent flux from diffusion, vesicle damage or detection background.",
                        "blocks": [{"block_type": "graph_data", "content": "A saturating curve supports carrier-limited transport, while a linear leak-like signal suggests uncontrolled diffusion or assay background."}],
                    },
                ],
            },
        ],
    }


def run_self_test() -> dict[str, Any]:
    failures: list[str] = []
    valid = sample_plan()
    valid_errors = validate_public_lecture_notes_plan(valid, "exam_prep_notes_docx")
    if valid_errors:
        failures.append("valid_plan_failed:" + ",".join(valid_errors))
    cases = {
        "course_map": {**valid, "course_knowledge_map": "Course Knowledge Map"},
        "chinese_default": {
            **valid,
            "public_lecture_sections": [
                {
                    "lecture_title": "Lecture 1",
                    "modules": [
                        {
                            "module_title": "膜结构",
                            "knowledge_functions": ["definition_boundary", "mechanism_process"],
                            "explanation": "这是一段中文默认输出，应该失败 because default public notes are English-only unless requested.",
                            "blocks": [{"block_type": "definition", "content": "Chinese public prose"}],
                        }
                    ],
                }
            ],
        },
        "flat_old_plan": {
            "object_type": "ExamPrepNotesPlan",
            "title": "Old plan",
            "course_modules": [],
            "source_scale_budget": valid["source_scale_budget"],
            "route_docx_style_profile": valid["route_docx_style_profile"],
        },
        "shallow": {
            **valid,
            "public_lecture_sections": [
                {
                    "lecture_title": "Lecture 1",
                    "modules": [
                        {
                            "module_title": "Too shallow",
                            "knowledge_functions": ["definition_boundary"],
                            "explanation": "This short module is intentionally too shallow.",
                            "blocks": [{"block_type": "definition", "content": "Short"}],
                        }
                    ],
                }
            ],
        },
        "strategy": {
            **valid,
            "public_lecture_sections": [
                {
                    "lecture_title": "Lecture 1",
                    "modules": [
                        {
                            "module_title": "Section A Strategy",
                            "knowledge_functions": ["definition_boundary", "mechanism_process"],
                            "explanation": "A strong answer should use this module for a recommended approach to exam strategy, which is not public knowledge prose.",
                            "blocks": [{"block_type": "explanation", "content": "Use this module."}],
                        }
                    ],
                }
            ],
        },
    }
    for name, plan in cases.items():
        if not validate_public_lecture_notes_plan(plan, "exam_prep_notes_docx"):
            failures.append(f"negative_case_passed:{name}")
    with tempfile.TemporaryDirectory(prefix="public_notes_selftest_") as tmp:
        out = Path(tmp) / "out"
        qa = Path(tmp) / "qa"
        manifest = write_public_lecture_notes_docx(valid, "exam_prep_notes_docx", out, qa, True, "manifest.json")
        if manifest.get("status") != "pass":
            failures.append("valid_render_failed")
        if not (out / "Lecture_Knowledge_Walkthrough.docx").exists():
            failures.append("valid_docx_missing")
    return {"pass": not failures, "failures": failures}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--plan", type=Path)
    parser.add_argument("--route", choices=sorted(ALLOWED_ROUTES), default="exam_prep_notes_docx")
    parser.add_argument("--output-dir", type=Path)
    parser.add_argument("--qa-dir", type=Path)
    parser.add_argument("--manifest-name", default="public_lecture_notes_manifest.json")
    parser.add_argument("--clean", action="store_true")
    parser.add_argument("--strict", action="store_true")
    parser.add_argument("--self-test", action="store_true")
    args = parser.parse_args()

    if args.self_test:
        result = run_self_test()
        print(json.dumps(result, indent=2))
        return 0 if result["pass"] else 1
    if not args.plan or not args.output_dir:
        print(json.dumps({"status": "fail", "qa_flags": ["missing_plan_or_output_dir"]}, indent=2))
        return 1
    qa_dir = args.qa_dir or args.output_dir
    if args.clean:
        if args.output_dir.exists():
            shutil.rmtree(args.output_dir)
        if qa_dir.exists() and qa_dir != args.output_dir:
            shutil.rmtree(qa_dir)
    result = write_public_lecture_notes_docx(load_plan(args.plan), args.route, args.output_dir, qa_dir, args.strict, args.manifest_name)
    print(json.dumps(result, indent=2))
    return 0 if result.get("status") in {"pass", "warn"} else 1


if __name__ == "__main__":
    raise SystemExit(main())
