#!/usr/bin/env python3
"""Lint public notes readability and layout segmentation."""

from __future__ import annotations

import argparse
import json
import re
import tempfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None  # type: ignore


WORD_RE = re.compile(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?")
FORMULA_RE = re.compile(r"(?:[A-Za-z]\s*=\s*[^.!?]{2,}|[A-Za-z]\s*/\s*[A-Za-z]|[A-Za-z]\s*\+\s*[A-Za-z])")
COLLAPSED_LIST_RE = re.compile(
    r"\b(?:include|includes|including|consist(?:s)?\s+of|criteria\s+are|steps\s+are|components\s+are)\b"
    r"[^.!?]{0,220}(?:,|;)[^.!?]{0,220}(?:,|;)[^.!?]{0,220}(?:,|;)",
    re.I,
)
GENERIC_TITLE = re.compile(r"^\s*(?:overview|introduction|background|summary|key\s+points|main\s+ideas|important\s+concepts)\s*$", re.I)
LOCAL_PATH_RE = re.compile(r"(^|[\s\"'(:])(?:/[^,\s]+|~[/\\][^,\s]+|[A-Za-z]:\\[^,\s]+|file://[^,\s]+)")


def words(text: str) -> list[str]:
    return WORD_RE.findall(str(text))


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def split_paragraphs(text: str) -> list[str]:
    return [part.strip() for part in re.split(r"\n{1,}|\r{1,}", str(text or "")) if part.strip()]


def block_text(block: Any) -> str:
    if not isinstance(block, dict):
        return str(block or "")
    content = block.get("content")
    if isinstance(content, list):
        return "\n".join(str(item).strip() for item in content if str(item).strip())
    return str(content or "").strip()


def iter_modules(plan: dict[str, Any]) -> list[tuple[str, dict[str, Any]]]:
    modules: list[tuple[str, dict[str, Any]]] = []
    for lecture_index, lecture in enumerate(plan.get("public_lecture_sections", []) or [], start=1):
        if not isinstance(lecture, dict):
            continue
        for module_index, module in enumerate(lecture.get("modules", []) or [], start=1):
            if isinstance(module, dict):
                modules.append((f"public_lecture_sections[{lecture_index}].modules[{module_index}]", module))
    return modules


def lint_plan(plan: dict[str, Any]) -> dict[str, Any]:
    failures: list[dict[str, Any]] = []
    module_count = 0
    for where, module in iter_modules(plan):
        module_count += 1
        title = str(module.get("module_title") or "").strip()
        explanation = str(module.get("explanation") or "").strip()
        blocks = module.get("blocks") if isinstance(module.get("blocks"), list) else []
        visual_refs = module.get("visual_refs") if isinstance(module.get("visual_refs"), list) else []
        if GENERIC_TITLE.match(title):
            failures.append({"type": "module_heading_not_micro_topic", "where": where, "module_title": title})
        if len(words(title)) > 16:
            failures.append({"type": "module_heading_too_long", "where": where, "module_title": title, "words": len(words(title))})
        for paragraph in split_paragraphs(explanation):
            count = len(words(paragraph))
            if count > 130:
                failures.append({"type": "paragraph_too_dense", "where": where, "module_title": title, "words": count})
            if FORMULA_RE.search(paragraph) and count > 45:
                failures.append({"type": "formula_or_equation_buried_in_long_prose", "where": where, "module_title": title})
            if COLLAPSED_LIST_RE.search(paragraph) and not any(isinstance(block.get("content"), list) for block in blocks if isinstance(block, dict)):
                failures.append({"type": "criteria_or_components_list_collapsed_into_prose", "where": where, "module_title": title})
        for block_index, block in enumerate(blocks, start=1):
            if not isinstance(block, dict):
                continue
            block_type = str(block.get("block_type") or "")
            content = block_text(block)
            for paragraph in split_paragraphs(content):
                count = len(words(paragraph))
                if count > 130 and block_type not in {"table", "calculation", "graph_data"}:
                    failures.append({"type": "block_paragraph_too_dense", "where": f"{where}.blocks[{block_index}]", "module_title": title, "words": count})
                if FORMULA_RE.search(paragraph) and count > 45 and block_type != "calculation":
                    failures.append({"type": "formula_not_separated_into_calculation_block", "where": f"{where}.blocks[{block_index}]", "module_title": title})
                if COLLAPSED_LIST_RE.search(paragraph) and not isinstance(block.get("content"), list):
                    failures.append({"type": "list_block_should_use_array_items", "where": f"{where}.blocks[{block_index}]", "module_title": title})
        for visual_ref in visual_refs:
            ref = str(visual_ref)
            if LOCAL_PATH_RE.search(ref):
                failures.append({"type": "visual_ref_contains_local_path", "where": where, "module_title": title})
            if not ref.strip():
                failures.append({"type": "empty_visual_ref", "where": where, "module_title": title})
        if visual_refs and not any(str(block.get("block_type") or "") in {"graph_data", "method", "calculation", "table", "explanation"} for block in blocks if isinstance(block, dict)):
            failures.append({"type": "visual_ref_without_nearby_explanation_block", "where": where, "module_title": title})
    if module_count and module_count < 2:
        failures.append({"type": "too_few_micro_modules_for_readability", "modules": module_count})
    return {"pass": not failures, "failures": failures}


def lint_docx(path: Path) -> dict[str, Any]:
    if Document is None:
        raise RuntimeError("python-docx is required to inspect DOCX files")
    doc = Document(path)
    failures: list[dict[str, Any]] = []
    visible_paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for index, text in enumerate(visible_paragraphs, start=1):
        count = len(words(text))
        if count > 150:
            failures.append({"type": "docx_visible_paragraph_too_dense", "paragraph": index, "words": count})
        if META_PRIVATE_PATH(text):
            failures.append({"type": "docx_contains_private_path", "paragraph": index})
    return {"pass": not failures, "paragraphs": len(visible_paragraphs), "failures": failures}


def META_PRIVATE_PATH(text: str) -> bool:
    return bool(LOCAL_PATH_RE.search(text))


def positive_plan() -> dict[str, Any]:
    return {
        "public_lecture_sections": [
            {
                "lecture_title": "Lecture 1",
                "modules": [
                    {
                        "module_title": "Initial slope gives initial reaction rate",
                        "knowledge_functions": ["method_readout", "graph_data_interpretation", "limitation_trap"],
                        "explanation": "The initial rate is the early linear change in signal per unit time. It is used because later curve sections can reflect substrate depletion or product effects, so the first straight portion gives the least distorted estimate of velocity.",
                        "blocks": [
                            {"block_type": "graph_data", "content": "Read the gradient from the first linear section of the graph."},
                            {"block_type": "limitation", "content": "Later curvature should not be used for initial velocity because reaction conditions have already drifted."},
                        ],
                    },
                    {
                        "module_title": "Dimension-aware substitution prevents concentration errors",
                        "knowledge_functions": ["calculation_unit_worked_example", "limitation_trap"],
                        "explanation": "A calculation is exam-ready only when the equation, units and substitution order are visible. Dimension-aware substitution prevents a correct formula from producing the wrong scale because each numerical value is converted before the final concentration is interpreted.",
                        "blocks": [
                            {"block_type": "calculation", "label": "Worked example", "content": "Use A = epsilon c l; rearrange for c; convert the answer into the requested unit."}
                        ],
                    },
                ],
            }
        ]
    }


def make_bad_docx(path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required to create DOCX files")
    doc = Document()
    doc.add_paragraph(" ".join(f"word{idx}" for idx in range(170)))
    doc.save(path)


def self_test() -> dict[str, Any]:
    good = lint_plan(positive_plan())
    bad = positive_plan()
    bad["public_lecture_sections"][0]["modules"][0]["explanation"] = (
        "The method includes sample preparation, calibration, blank correction, absorbance measurement, unit conversion, graphing, rate calculation, comparison, interpretation and reporting."
    )
    collapsed = lint_plan(bad)
    dense = positive_plan()
    dense["public_lecture_sections"][0]["modules"][0]["explanation"] = " ".join(f"word{idx}" for idx in range(150))
    dense_result = lint_plan(dense)
    with tempfile.TemporaryDirectory(prefix="notes_readability_linter_") as tmp:
        bad_docx = Path(tmp) / "bad.docx"
        make_bad_docx(bad_docx)
        bad_docx_result = lint_docx(bad_docx)
    failures: list[dict[str, Any]] = []
    if not good["pass"]:
        failures.append({"type": "positive_plan_rejected", "result": good})
    if collapsed["pass"]:
        failures.append({"type": "collapsed_list_not_rejected", "result": collapsed})
    if dense_result["pass"]:
        failures.append({"type": "dense_paragraph_not_rejected", "result": dense_result})
    if bad_docx_result["pass"]:
        failures.append({"type": "dense_docx_not_rejected", "result": bad_docx_result})
    return {"pass": not failures, "positive": good, "collapsed_negative": collapsed, "dense_negative": dense_result, "docx_negative": bad_docx_result, "failures": failures}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--plan", type=Path)
    parser.add_argument("--docx", type=Path)
    parser.add_argument("--self-test", action="store_true")
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    if args.self_test:
        result = self_test()
    elif args.plan:
        result = lint_plan(load_json(args.plan))
    elif args.docx:
        result = lint_docx(args.docx)
    else:
        result = {"pass": False, "failures": [{"type": "missing_plan_docx_or_self_test"}]}
    text = json.dumps(result, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text + "\n", encoding="utf-8")
    else:
        print(text)
    return 0 if result.get("pass") else 1


if __name__ == "__main__":
    raise SystemExit(main())
