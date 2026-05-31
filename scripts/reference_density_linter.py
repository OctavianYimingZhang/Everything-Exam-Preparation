#!/usr/bin/env python3
"""Compare public DOCX density against a reference-quality DOCX."""

from __future__ import annotations

import argparse
import json
import tempfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")


def count_words(text: str) -> int:
    return len([token for token in text.replace("\n", " ").split(" ") if token.strip()])


def docx_metrics(path: Path) -> dict[str, Any]:
    doc = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    body_paragraphs = [text for text in paragraphs if count_words(text) > 12]
    visible_words = sum(count_words(text) for text in paragraphs)
    body_words = sum(count_words(text) for text in body_paragraphs)
    return {
        "path": str(path),
        "paragraphs": len(paragraphs),
        "body_paragraphs": len(body_paragraphs),
        "heading_like_paragraphs": len([text for text in paragraphs if count_words(text) <= 12 and not text.endswith(".")]),
        "visible_words": visible_words,
        "body_words": body_words,
        "avg_body_words": round(body_words / max(1, len(body_paragraphs)), 1),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "bytes": path.stat().st_size,
    }


def lint(
    target: Path,
    *,
    reference: Path | None = None,
    min_words: int = 0,
    min_body_paragraphs: int = 0,
    min_reference_ratio: float = 0.0,
) -> dict[str, Any]:
    target_metrics = docx_metrics(target)
    reference_metrics = docx_metrics(reference) if reference else None
    failures: list[dict[str, Any]] = []
    if min_words and target_metrics["visible_words"] < min_words:
        failures.append({"type": "visible_words_below_minimum", "visible_words": target_metrics["visible_words"], "minimum": min_words})
    if min_body_paragraphs and target_metrics["body_paragraphs"] < min_body_paragraphs:
        failures.append(
            {
                "type": "body_paragraphs_below_minimum",
                "body_paragraphs": target_metrics["body_paragraphs"],
                "minimum": min_body_paragraphs,
            }
        )
    if reference_metrics and min_reference_ratio > 0:
        reference_floor = int(reference_metrics["visible_words"] * min_reference_ratio)
        if target_metrics["visible_words"] < reference_floor:
            failures.append(
                {
                    "type": "visible_words_below_reference_ratio",
                    "visible_words": target_metrics["visible_words"],
                    "reference_visible_words": reference_metrics["visible_words"],
                    "minimum_ratio": min_reference_ratio,
                    "minimum_visible_words": reference_floor,
                }
            )
        body_floor = max(1, int(reference_metrics["body_paragraphs"] * min_reference_ratio))
        if target_metrics["body_paragraphs"] < body_floor:
            failures.append(
                {
                    "type": "body_paragraphs_below_reference_ratio",
                    "body_paragraphs": target_metrics["body_paragraphs"],
                    "reference_body_paragraphs": reference_metrics["body_paragraphs"],
                    "minimum_ratio": min_reference_ratio,
                    "minimum_body_paragraphs": body_floor,
                }
            )
    return {"pass": not failures, "target": target_metrics, "reference": reference_metrics, "failures": failures}


def make_docx(path: Path, paragraphs: int, words_per_paragraph: int) -> None:
    doc = Document()
    for idx in range(paragraphs):
        text = " ".join([f"word{idx}_{j}" for j in range(words_per_paragraph)])
        doc.add_paragraph(text)
    doc.save(path)


def self_test() -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="reference_density_linter_") as tmp:
        tmp_path = Path(tmp)
        reference = tmp_path / "reference.docx"
        short = tmp_path / "short.docx"
        dense = tmp_path / "dense.docx"
        make_docx(reference, 72, 70)
        make_docx(short, 14, 50)
        make_docx(dense, 75, 72)
        bad = lint(short, reference=reference, min_reference_ratio=1.0)
        good = lint(dense, reference=reference, min_reference_ratio=1.0)
    failures: list[dict[str, Any]] = []
    if bad["pass"]:
        failures.append({"type": "short_docx_not_rejected", "result": bad})
    if not good["pass"]:
        failures.append({"type": "dense_docx_rejected", "result": good})
    return {"pass": not failures, "bad_result": bad, "good_result": good, "failures": failures}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("target", nargs="?", type=Path)
    parser.add_argument("--reference", type=Path)
    parser.add_argument("--min-words", type=int, default=0)
    parser.add_argument("--min-body-paragraphs", type=int, default=0)
    parser.add_argument("--min-reference-ratio", type=float, default=0.0)
    parser.add_argument("--self-test", action="store_true")
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    if args.self_test:
        result = self_test()
    elif args.target:
        result = lint(
            args.target,
            reference=args.reference,
            min_words=args.min_words,
            min_body_paragraphs=args.min_body_paragraphs,
            min_reference_ratio=args.min_reference_ratio,
        )
    else:
        result = {"pass": False, "failures": [{"type": "missing_target_or_self_test"}]}
    text = json.dumps(result, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text + "\n", encoding="utf-8")
    else:
        print(text)
    return 0 if result.get("pass") else 1


if __name__ == "__main__":
    raise SystemExit(main())
