#!/usr/bin/env python3
"""Lint Academic Exam-Ready Notes DOCX style."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Cm, RGBColor  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

BLUE_RGB = {"0000FF", "0563C1", "2F5496", "1F4E79", "4472C4", "5B9BD5"}
EXPECTED_MARGIN_CM = 2.0
MIN_LINE_SPACING = 1.45
MAX_LINE_SPACING = 1.55
FORBIDDEN_INTERNAL_HEADINGS = {
    "Exam Specificity",
    "Core Exam Claim",
    "Exam Use",
    "Common Error / Trap",
    "Must Master",
    "Course-Level Exam Map",
    "How To Answer This Exam",
}
COLOUR_RE = re.compile(rb"<w:color\b[^>]*/?>")
VAL_RE = re.compile(rb'w:val="([^"]+)"')
THEME_RE = re.compile(rb'w:themeColor="([^"]+)"')


def iter_docx_paths(path: Path) -> list[Path]:
    if path.is_file() and path.suffix.lower() == ".docx":
        return [path]
    if path.is_dir():
        return sorted(child for child in path.rglob("*.docx") if child.is_file())
    return []


def line_spacing_value(paragraph: Any) -> float | None:
    value = paragraph.paragraph_format.line_spacing
    if value is None and paragraph.style is not None:
        value = paragraph.style.paragraph_format.line_spacing
    if value is None:
        return None
    try:
        return float(value)
    except TypeError:
        return None


def paragraph_has_image(paragraph: Any) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def is_forbidden_heading(text: str) -> bool:
    return text.strip().rstrip(":").casefold() in {heading.casefold() for heading in FORBIDDEN_INTERNAL_HEADINGS}


def is_heading(paragraph: Any, visible_index: int) -> bool:
    style_name = (paragraph.style.name if paragraph.style else "").casefold()
    text = paragraph.text.strip()
    if visible_index == 1:
        return True
    if any(key in style_name for key in ["title", "heading", "subheading", "lecture", "module"]):
        return True
    if text.startswith("Lecture:") or text.startswith("Module:"):
        return True
    return len(text) <= 130 and text and not text.endswith(".") and any(run.bold for run in paragraph.runs)


def colour_string(colour: Any) -> str | None:
    if colour is None:
        return None
    if getattr(colour, "rgb", None):
        return str(colour.rgb).upper()
    if getattr(colour, "theme_color", None):
        return f"theme:{colour.theme_color}"
    return None


def lint_docx(path: Path) -> list[dict[str, Any]]:
    failures: list[dict[str, Any]] = []
    doc = Document(path)
    if not doc.sections:
        return [{"type": "missing_docx_section", "path": str(path)}]
    section = doc.sections[0]
    for margin_name in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        value = float(getattr(section, margin_name).cm)
        if abs(value - EXPECTED_MARGIN_CM) > 0.08:
            failures.append({"type": "bad_margin", "path": str(path), "margin": margin_name, "cm": round(value, 3)})

    visible_index = 0
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        has_image = paragraph_has_image(paragraph)
        if not text and not has_image:
            continue
        visible_index += 1
        if text and is_forbidden_heading(text):
            failures.append({"type": "forbidden_internal_heading", "path": str(path), "paragraph": index, "text": text})
        spacing = line_spacing_value(paragraph)
        heading = is_heading(paragraph, visible_index)
        if has_image:
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                failures.append({"type": "image_not_centered", "path": str(path), "paragraph": index})
        elif heading:
            if paragraph.alignment not in {None, WD_ALIGN_PARAGRAPH.LEFT}:
                failures.append({"type": "heading_not_left_aligned", "path": str(path), "paragraph": index})
        else:
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                failures.append({"type": "body_not_justified", "path": str(path), "paragraph": index, "alignment": str(paragraph.alignment)})
        if spacing is not None and not (MIN_LINE_SPACING <= spacing <= MAX_LINE_SPACING) and not has_image:
            failures.append({"type": "bad_line_spacing", "path": str(path), "paragraph": index, "line_spacing": round(spacing, 3)})

        style_colour = colour_string(paragraph.style.font.color if paragraph.style and paragraph.style.font else None)
        if style_colour and style_colour != "000000":
            failures.append({"type": "style_font_colour_not_black", "path": str(path), "paragraph": index, "colour": style_colour})
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            colour = colour_string(run.font.color)
            if colour is not None and colour != "000000":
                failures.append({"type": "run_font_colour_not_black", "path": str(path), "paragraph": index, "text": run.text[:80], "colour": colour})
            if colour in BLUE_RGB or (colour or "").startswith("theme:"):
                failures.append({"type": "blue_or_theme_text_detected", "path": str(path), "paragraph": index, "text": run.text[:80], "colour": colour})
            font_names = {name for name in [run.font.name, paragraph.style.font.name if paragraph.style else None] if name}
            if font_names and "Arial" not in font_names:
                failures.append({"type": "non_arial_text", "path": str(path), "paragraph": index, "fonts": sorted(font_names)})
    try:
        with zipfile.ZipFile(path) as archive:
            name = "word/document.xml"
            if name in archive.namelist():
                xml = archive.read(name)
                for match in COLOUR_RE.finditer(xml):
                    tag = match.group(0)
                    value_match = VAL_RE.search(tag)
                    theme_match = THEME_RE.search(tag)
                    value = (value_match.group(1) if value_match else b"").decode("ascii", errors="ignore").upper()
                    theme = (theme_match.group(1) if theme_match else b"").decode("ascii", errors="ignore")
                    if theme or (value and value not in {"000000", "AUTO"}):
                        failures.append({"type": "ooxml_non_black_or_theme_colour", "path": str(path), "part": name, "value": value, "theme": theme})
    except Exception as exc:
        failures.append({"type": "ooxml_colour_scan_error", "path": str(path), "error": str(exc)})
    return failures


def lint_path(path: Path) -> dict[str, Any]:
    docx_paths = iter_docx_paths(path)
    failures: list[dict[str, Any]] = []
    if not docx_paths:
        failures.append({"type": "no_docx_found", "path": str(path)})
    for docx_path in docx_paths:
        failures.extend(lint_docx(docx_path))
    return {"pass": not failures, "counts": {"docx_files": len(docx_paths)}, "failures": failures}


def create_bad_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    paragraph = doc.add_paragraph("Exam Use:")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("Blue text must fail.")
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 0, 255)
    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("path", type=Path, nargs="?")
    parser.add_argument("--self-test-bad", action="store_true")
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    tmp_dir: str | None = None
    try:
        if args.self_test_bad:
            tmp_dir = tempfile.mkdtemp(prefix="exam_prep_style_bad_")
            target = Path(tmp_dir) / "bad_style.docx"
            create_bad_docx(target)
            result = lint_path(target)
        elif args.path:
            result = lint_path(args.path)
        else:
            result = {"pass": False, "failures": [{"type": "missing_path"}], "counts": {"docx_files": 0}}
    except Exception as exc:
        result = {"pass": False, "failures": [{"type": "read_error", "error": str(exc)}], "counts": {"docx_files": 0}}
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)
    text = json.dumps(result, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text + "\n", encoding="utf-8")
    else:
        print(text)
    return 0 if result.get("pass") else 1


if __name__ == "__main__":
    raise SystemExit(main())
