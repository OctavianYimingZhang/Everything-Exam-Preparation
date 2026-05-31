#!/usr/bin/env python3
"""Generate the default Academic Exam-Ready Notes DOCX from conceptual modules.

This route intentionally does not render raw slide inventories, extracted slide
bullets, administrative content, or internal scaffolds. The plan must provide a
conceptual course map, source-scale budget and examinable knowledge units.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Cm, Pt, RGBColor  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

from knowledge_only_rendering_rules import (
    forbidden_advisory_heading_hits,
    forbidden_advisory_phrase_hits,
    forbidden_non_knowledge_hits,
    repeated_template_label_hits,
)

STYLE_DEFAULTS = {
    "route": "exam_prep_notes_docx",
    "margin_cm": 2.0,
    "line_spacing": 1.5,
    "body_alignment": "justified",
    "title_alignment": "left",
    "heading_alignment": "left",
    "image_alignment": "center",
    "body_font_pt": 10.5,
    "heading_font_pt": 12.0,
    "module_heading_font_pt": 14.0,
    "title_font_pt": 14.0,
    "text_color": "black",
    "module_page_breaks": False,
    "theme_colours_allowed": False,
    "blue_heading_styles_allowed": False,
}

FORBIDDEN_TEXT = {
    "source anchor",
    "confidence",
    "evidence score",
    "recurrence count",
    "examiner operation",
    "discriminator axis",
    "exam specificity",
    "core exam claim",
    "exam use",
    "common error / trap",
    "must master",
    "practice question",
    "answer key",
    "past paper year",
    "prediction score",
    "according to slide",
    "english explanations extracted",
    "ppt page",
    "slide mentions",
    "slides say",
    "the final slide",
    "the first slide",
    "the next slide",
    "the second slide",
    "this slide",
}


def load_plan(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def style_value(plan: dict[str, Any], key: str) -> Any:
    profile = plan.get("route_docx_style_profile")
    if isinstance(profile, dict) and key in profile:
        return profile[key]
    return STYLE_DEFAULTS[key]


def normalize_run(run, size_pt: float | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size_pt or float(STYLE_DEFAULTS["body_font_pt"]))
    run.font.color.rgb = RGBColor(0, 0, 0)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def paragraph_size(plan: dict[str, Any], kind: str) -> float:
    return {
        "title": float(style_value(plan, "title_font_pt")),
        "module": float(style_value(plan, "module_heading_font_pt")),
        "heading": float(style_value(plan, "heading_font_pt")),
        "subheading": float(style_value(plan, "heading_font_pt")),
        "body": float(style_value(plan, "body_font_pt")),
    }[kind]


def normalize_paragraph(paragraph, kind: str, plan: dict[str, Any]) -> None:
    paragraph.paragraph_format.line_spacing = float(style_value(plan, "line_spacing"))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if kind == "body" else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        normalize_run(run, paragraph_size(plan, kind))


def set_document_defaults(doc: Document, plan: dict[str, Any]) -> None:
    margin_cm = float(style_value(plan, "margin_cm"))
    section = doc.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    for style_name in ["Normal", "EPNTitle", "EPNModule", "EPNHeading", "EPNSubheading", "EPNBody"]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, 1)
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(float(style_value(plan, "body_font_pt")))
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = float(style_value(plan, "line_spacing"))
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)


def add_marked_text(paragraph, text: str, size_pt: float | None = None) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*)", str(text)):
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
        normalize_run(run, size_pt)


def add_paragraph(doc: Document, text: str, plan: dict[str, Any], kind: str = "body"):
    if not str(text or "").strip():
        return None
    style = {"title": "EPNTitle", "module": "EPNModule", "heading": "EPNHeading", "subheading": "EPNSubheading", "body": "EPNBody"}[kind]
    paragraph = doc.add_paragraph(style=style)
    add_marked_text(paragraph, str(text).strip(), paragraph_size(plan, kind))
    if kind in {"title", "module", "heading", "subheading"}:
        for run in paragraph.runs:
            run.bold = True
    normalize_paragraph(paragraph, kind, plan)
    return paragraph


def get_course_modules(plan: dict[str, Any]) -> list[dict[str, Any]]:
    return [module for module in plan.get("course_modules", []) or [] if isinstance(module, dict)]


def visible_strings(plan: dict[str, Any]) -> list[str]:
    parts: list[str] = [str(plan.get("title") or ""), str(plan.get("course_knowledge_map") or "")]
    for module in get_course_modules(plan):
        parts.extend(str(module.get(key) or "") for key in ["module_title", "module_function"])
        for unit in module.get("examinable_units", []) or []:
            if isinstance(unit, dict):
                parts.extend(str(unit.get(key) or "") for key in ["title", "explanation", "optional_equation_or_example", "common_confusion_or_boundary"])
    return [part for part in parts if part.strip()]


def star_prefix(priority: Any) -> str:
    mapping = {"high": "★★★", "medium": "★★", "low": "★", "★★★": "★★★", "★★": "★★", "★": "★"}
    value = str(priority or "medium").strip().lower()
    return mapping.get(value, mapping.get(str(priority or "").strip(), "★★"))


def infer_source_units(plan: dict[str, Any]) -> int:
    budget = plan.get("source_scale_budget")
    if isinstance(budget, dict):
        for key in ["source_units_count", "readable_source_blocks", "protected_knowledge_units_total"]:
            value = budget.get(key)
            if isinstance(value, int) and value > 0:
                return value
    seen: set[str] = set()
    for module in get_course_modules(plan):
        for source in module.get("source_lectures", []) or []:
            if str(source).strip():
                seen.add(str(source).strip())
    return len(seen)


def count_public_units(plan: dict[str, Any]) -> int:
    return sum(len([unit for unit in module.get("examinable_units", []) or [] if isinstance(unit, dict)]) for module in get_course_modules(plan))


def floor_for_source_units(source_units: int) -> int:
    if source_units <= 0:
        return 0
    if source_units <= 3:
        return 8
    if source_units <= 8:
        return 20
    if source_units <= 15:
        return 40
    return 70


def validate_source_scale_budget(plan: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    budget = plan.get("source_scale_budget")
    if not isinstance(budget, dict):
        return ["exam_prep_notes_requires_source_scale_budget"]
    source_units = infer_source_units(plan)
    public_units = count_public_units(plan)
    target_min = budget.get("target_public_units_min")
    if not isinstance(target_min, int) or target_min < 0:
        errors.append("source_scale_budget_missing_target_public_units_min")
        target_min = floor_for_source_units(source_units)
    else:
        target_min = max(target_min, floor_for_source_units(source_units))
    if budget.get("coverage_floor_status") == "block":
        errors.append("source_scale_budget_blocks_generation")
    if public_units < int(target_min):
        errors.append(f"source_scale_budget_public_units_too_low:{public_units}<{int(target_min)}")
    if budget.get("compression_mode") not in {"explain_not_dump", "concise", "standard", "expanded", "multi_volume_required"}:
        errors.append("source_scale_budget_bad_compression_mode")
    return errors


def validate_public_density(plan: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    short_units = []
    for module in get_course_modules(plan):
        for unit in module.get("examinable_units", []) or []:
            if not isinstance(unit, dict):
                continue
            explanation = str(unit.get("explanation") or "").strip()
            if len(explanation.split()) < 35:
                short_units.append(str(unit.get("title") or "Knowledge unit"))
    if short_units:
        errors.append("examinable_units_too_short_for_reference_caliber:" + "; ".join(short_units[:5]))
    combined = "\n".join(visible_strings(plan))
    word_count = len([token for token in combined.replace("\n", " ").split(" ") if token.strip()])
    budget = plan.get("source_scale_budget")
    target_words_min = budget.get("target_words_min") if isinstance(budget, dict) else None
    if isinstance(target_words_min, int) and target_words_min > 0 and word_count < target_words_min:
        errors.append(f"source_scale_budget_visible_words_too_low:{word_count}<{target_words_min}")
    return errors


def validate_plan(plan: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if plan.get("object_type") != "ExamPrepNotesPlan":
        errors.append("plan_object_type_not_exam_prep_notes_plan")
    if not get_course_modules(plan):
        errors.append("exam_prep_notes_requires_course_modules")
    errors.extend(validate_source_scale_budget(plan))
    errors.extend(validate_public_density(plan))
    profile = plan.get("route_docx_style_profile")
    if not isinstance(profile, dict):
        errors.append("exam_prep_notes_requires_route_docx_style_profile")
        profile = {}
    if profile.get("route") != "exam_prep_notes_docx":
        errors.append("exam_prep_notes_style_profile_wrong_route")
    if profile.get("body_alignment") != "justified":
        errors.append("exam_prep_notes_style_profile_body_not_justified")
    if profile.get("title_alignment") != "left":
        errors.append("exam_prep_notes_style_profile_title_not_left")
    if profile.get("heading_alignment") != "left":
        errors.append("exam_prep_notes_style_profile_heading_not_left")
    if profile.get("image_alignment") != "center":
        errors.append("exam_prep_notes_style_profile_image_not_center")
    if profile.get("text_color") != "black":
        errors.append("exam_prep_notes_style_profile_text_not_black")
    if profile.get("theme_colours_allowed") is not False:
        errors.append("exam_prep_notes_theme_colours_not_disabled")
    if profile.get("blue_heading_styles_allowed") is not False:
        errors.append("exam_prep_notes_blue_heading_styles_not_disabled")
    spacing = profile.get("line_spacing")
    if not isinstance(spacing, (int, float)) or not (1.45 <= float(spacing) <= 1.55):
        errors.append("exam_prep_notes_style_profile_line_spacing_not_1_5")
    for value in visible_strings(plan):
        lowered = value.lower()
        for phrase in FORBIDDEN_TEXT:
            if phrase in lowered:
                errors.append(f"forbidden_text_in_plan:{phrase}")
        for phrase in forbidden_advisory_phrase_hits(value):
            errors.append(f"forbidden_advisory_phrase_in_plan:{phrase}")
        for heading in forbidden_advisory_heading_hits(value):
            errors.append(f"forbidden_advisory_heading_in_plan:{heading}")
        for category in forbidden_non_knowledge_hits(value):
            errors.append(f"forbidden_non_knowledge_surface_in_plan:{category}")
    combined_visible_text = "\n".join(visible_strings(plan))
    for label in repeated_template_label_hits(combined_visible_text):
        errors.append(f"repeated_rigid_template_label:{label}")
    return sorted(set(errors))


def write_docx(plan: dict[str, Any], output_dir: Path, qa_dir: Path, strict: bool) -> dict[str, Any]:
    errors = validate_plan(plan)
    if strict and errors:
        return {"status": "fail", "qa_flags": errors, "documents": []}

    doc = Document()
    set_document_defaults(doc, plan)
    title = str(plan.get("title") or "Academic Exam-Ready Knowledge Notes")
    add_paragraph(doc, title, plan, "title")
    if plan.get("course_knowledge_map"):
        add_paragraph(doc, "Course Knowledge Map", plan, "heading")
        add_paragraph(doc, str(plan["course_knowledge_map"]), plan, "body")

    manifest = {
        "notes_plan_id": plan.get("notes_plan_id"),
        "target_group_key": plan.get("target_group_key"),
        "title": title,
        "source_scale_budget": plan.get("source_scale_budget"),
        "public_units": count_public_units(plan),
        "modules": [],
        "qa_flags": errors,
    }
    for idx, module in enumerate(get_course_modules(plan)):
        if idx > 0 and bool(style_value(plan, "module_page_breaks")):
            doc.add_page_break()
        module_title = str(module.get("module_title") or "Course module").strip()
        add_paragraph(doc, module_title, plan, "module")
        add_paragraph(doc, str(module.get("module_function") or ""), plan, "body")
        core_questions = module.get("core_questions") or []
        if core_questions:
            add_paragraph(doc, "Core questions: " + "; ".join(str(item).strip() for item in core_questions if str(item).strip()), plan, "body")
        module_manifest = {"module_title": module_title, "units": []}
        for unit in module.get("examinable_units", []) or []:
            if not isinstance(unit, dict):
                continue
            unit_title = str(unit.get("title") or "Knowledge unit").strip()
            add_paragraph(doc, f"{star_prefix(unit.get('priority'))} {unit_title}", plan, "subheading")
            add_paragraph(doc, str(unit.get("explanation") or ""), plan, "body")
            add_paragraph(doc, str(unit.get("optional_equation_or_example") or ""), plan, "body")
            add_paragraph(doc, str(unit.get("common_confusion_or_boundary") or ""), plan, "body")
            module_manifest["units"].append({"title": unit_title, "priority": unit.get("priority")})
        manifest["modules"].append(module_manifest)

    output_dir.mkdir(parents=True, exist_ok=True)
    qa_dir.mkdir(parents=True, exist_ok=True)
    filename = "Lecture_Knowledge_Walkthrough.docx"
    docx_path = output_dir / filename
    doc.save(docx_path)
    manifest["documents"] = [{"docx_path": str(docx_path), "filename": filename}]
    manifest["route_docx_style_profile"] = plan.get("route_docx_style_profile") or STYLE_DEFAULTS
    manifest["status"] = "pass" if not errors else "warn"
    (qa_dir / "exam_prep_notes_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--plan", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--qa-dir", type=Path)
    parser.add_argument("--clean", action="store_true")
    parser.add_argument("--strict", action="store_true")
    parser.add_argument("--deliverable-only", action="store_true")
    args = parser.parse_args()

    output_dir = args.output_dir
    qa_dir = args.qa_dir or (output_dir if not args.deliverable_only else output_dir.parent / "exam_prep_notes_internal_qa")
    if args.clean:
        if output_dir.exists():
            shutil.rmtree(output_dir)
        if qa_dir.exists() and qa_dir != output_dir:
            shutil.rmtree(qa_dir)

    result = write_docx(load_plan(args.plan), output_dir, qa_dir, args.strict)
    print(json.dumps(result, indent=2))
    return 0 if result.get("status") in {"pass", "warn"} else 1


if __name__ == "__main__":
    raise SystemExit(main())
