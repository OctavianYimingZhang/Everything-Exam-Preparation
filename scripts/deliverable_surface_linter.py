#!/usr/bin/env python3
"""Route-based deliverable surface QA for public notes and Example Essay DOCX."""

from __future__ import annotations

import argparse
import json
import tempfile
import zipfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Cm, Pt  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

HELPER_NAMES = {
    "example_essay_manifest.json",
    "example_essay_source_audit.json",
    "citation_candidates.json",
    "citation_resolution_log.json",
    "citation_source_notes.json",
    "classic_experiment_search_plan.json",
    "source_scan.json",
    "target_groups.json",
}
HELPER_SUFFIXES = ("_source_map.json", "_qa.json", "_render_qa.json")
FORBIDDEN_PUBLIC_SUFFIXES = {".xlsx", ".xlsm"}
PUBLIC_NOTES_FORBIDDEN = ("course knowledge map", "source role summary", "this slide", "workflow", "qa flag")
CM_TOL = 0.08


def emu_to_cm(value: Any) -> float:
    return float(value.cm) if value is not None else 0.0


def close(value: float, expected: float, tol: float = CM_TOL) -> bool:
    return abs(value - expected) <= tol


def collect_docx(path: Path) -> list[Path]:
    if path.is_dir():
        return sorted(path.rglob("*.docx"))
    return [path] if path.suffix.lower() == ".docx" else []


def lint_public_folder(path: Path, allowed: set[str] | None = None) -> list[dict[str, Any]]:
    failures: list[dict[str, Any]] = []
    files = sorted(item for item in path.rglob("*") if item.is_file()) if path.is_dir() else [path]
    for item in files:
        name = item.name
        if name in HELPER_NAMES or any(name.endswith(suffix) for suffix in HELPER_SUFFIXES):
            failures.append({"type": "helper_artifact_in_public_output", "path": str(item)})
        elif item.suffix.lower() in FORBIDDEN_PUBLIC_SUFFIXES:
            failures.append({"type": "legacy_workbook_in_public_output", "path": str(item)})
        elif allowed and item.suffix.lower() not in allowed:
            failures.append({"type": "non_deliverable_file", "path": str(item), "suffix": item.suffix})
    return failures


def lint_public_notes_docx(path: Path) -> list[dict[str, Any]]:
    doc = Document(path)
    failures: list[dict[str, Any]] = []
    section = doc.sections[0]
    margins = [section.top_margin, section.bottom_margin, section.left_margin, section.right_margin]
    if not all(close(emu_to_cm(value), 2.0) for value in margins):
        failures.append({"type": "public_notes_margins_not_2_0_cm", "path": str(path)})
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        lower = text.lower()
        for phrase in PUBLIC_NOTES_FORBIDDEN:
            if phrase in lower:
                failures.append({"type": "forbidden_public_notes_surface", "path": str(path), "paragraph": index, "phrase": phrase})
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and index > 1:
            failures.append({"type": "unexpected_centered_public_notes_paragraph", "path": str(path), "paragraph": index})
        for run_index, run in enumerate(paragraph.runs, start=1):
            font_name = run.font.name
            if font_name and font_name.lower() != "arial":
                failures.append({"type": "non_arial_text", "path": str(path), "paragraph": index, "run": run_index, "font": font_name})
    return failures


def load_source_map(docx_path: Path, explicit_map: Path | None = None) -> dict[str, Any] | None:
    if explicit_map:
        return json.loads(explicit_map.read_text(encoding="utf-8"))
    for candidate in [docx_path.with_name(docx_path.stem.split("_")[0] + "_source_map.json"), docx_path.with_name(docx_path.stem + "_source_map.json")]:
        if candidate.exists():
            return json.loads(candidate.read_text(encoding="utf-8"))
    return None


def lint_example_essay_docx(path: Path, source_map_path: Path | None = None) -> list[dict[str, Any]]:
    doc = Document(path)
    source_map = load_source_map(path, source_map_path)
    failures: list[dict[str, Any]] = []
    section = doc.sections[0]
    margins = [section.top_margin, section.bottom_margin, section.left_margin, section.right_margin]
    if not all(close(emu_to_cm(value), 2.5) for value in margins):
        failures.append({"type": "essay_margins_not_2_5_cm", "path": str(path)})
    visible_index = 0
    source_paragraphs = source_map.get("paragraphs", []) if source_map else []
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            failures.append({"type": "empty_spacer_paragraph", "path": str(path)})
            continue
        visible_index += 1
        kind = source_paragraphs[visible_index - 1].get("kind") if visible_index - 1 < len(source_paragraphs) else ("title" if visible_index == 1 else "body")
        if kind == "title" and paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            failures.append({"type": "essay_title_not_centered", "path": str(path), "paragraph": visible_index})
        if kind == "body" and paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            failures.append({"type": "essay_body_not_justified", "path": str(path), "paragraph": visible_index})
        if paragraph.paragraph_format.line_spacing != 1.5:
            failures.append({"type": "essay_line_spacing_not_1_5", "path": str(path), "paragraph": visible_index})
        for run_index, run in enumerate(paragraph.runs, start=1):
            font_name = run.font.name
            if font_name and font_name.lower() != "arial":
                failures.append({"type": "essay_non_arial_text", "path": str(path), "paragraph": visible_index, "run": run_index, "font": font_name})
    if source_map:
        for index, paragraph in enumerate(source_paragraphs, start=1):
            if paragraph.get("kind") == "body" and not paragraph.get("lecture_anchors"):
                failures.append({"type": "body_paragraph_missing_lecture_anchor", "path": str(path), "paragraph": index})
            for run in paragraph.get("runs", []) or []:
                if run.get("highlight") == "green" and not run.get("in_text_citation"):
                    failures.append({"type": "green_run_missing_citation", "path": str(path), "paragraph": index})
                if run.get("highlight") == "yellow" and not run.get("source_anchor"):
                    failures.append({"type": "yellow_run_missing_extra_reading_anchor", "path": str(path), "paragraph": index})
    return failures


def structural_docx_check(path: Path) -> list[dict[str, Any]]:
    failures: list[dict[str, Any]] = []
    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
        for required in ["word/document.xml", "[Content_Types].xml"]:
            if required not in names:
                failures.append({"type": "docx_missing_part", "path": str(path), "part": required})
    except Exception as exc:
        failures.append({"type": "docx_structural_read_error", "path": str(path), "error": str(exc)})
    return failures


def lint_path(path: Path, route: str, source_map: Path | None = None, allowed: set[str] | None = None) -> dict[str, Any]:
    failures = lint_public_folder(path, allowed)
    for docx_path in collect_docx(path):
        failures.extend(structural_docx_check(docx_path))
        if route == "public_notes":
            failures.extend(lint_public_notes_docx(docx_path))
        elif route == "example_essay":
            failures.extend(lint_example_essay_docx(docx_path, source_map))
    return {"pass": not failures, "route": route, "path": str(path), "failures": failures}


def make_public_docx(path: Path, bad: bool = False) -> None:
    doc = Document()
    section = doc.sections[0]
    margin = Cm(2.5 if bad else 2.0)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = margin
    p = doc.add_paragraph("Course Knowledge Map: this slide exposes workflow." if bad else "Enzyme inhibition changes reaction rate because active-site access controls catalysis.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bad else WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = "Times New Roman" if bad else "Arial"
    doc.save(path)


def make_essay_docx(path: Path, bad: bool = False) -> None:
    doc = Document()
    section = doc.sections[0]
    margin = Cm(2.0 if bad else 2.5)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = margin
    p = doc.add_paragraph("Essay title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if bad else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0 if bad else 1.5
    p.runs[0].font.name = "Times New Roman" if bad else "Arial"
    body = doc.add_paragraph("Competitive inhibition raises the substrate requirement because active-site competition changes apparent affinity.")
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT if bad else WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.line_spacing = 1.0 if bad else 1.5
    body.runs[0].font.name = "Times New Roman" if bad else "Arial"
    doc.save(path)


def self_test() -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="deliverable_surface_") as tmp:
        root = Path(tmp)
        good_public = root / "good_public.docx"
        bad_public = root / "bad_public.docx"
        good_essay = root / "good_essay.docx"
        bad_essay = root / "bad_essay.docx"
        make_public_docx(good_public, bad=False)
        make_public_docx(bad_public, bad=True)
        make_essay_docx(good_essay, bad=False)
        make_essay_docx(bad_essay, bad=True)
        public_good = lint_path(good_public, "public_notes")
        public_bad = lint_path(bad_public, "public_notes")
        essay_good = lint_path(good_essay, "example_essay")
        essay_bad = lint_path(bad_essay, "example_essay")
    failures = []
    if not public_good["pass"]:
        failures.append({"type": "public_good_rejected", "result": public_good})
    if public_bad["pass"]:
        failures.append({"type": "public_bad_accepted", "result": public_bad})
    if not essay_good["pass"]:
        failures.append({"type": "essay_good_rejected", "result": essay_good})
    if essay_bad["pass"]:
        failures.append({"type": "essay_bad_accepted", "result": essay_bad})
    return {"pass": not failures, "public_good": public_good, "public_bad": public_bad, "essay_good": essay_good, "essay_bad": essay_bad, "failures": failures}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("path", type=Path, nargs="?")
    parser.add_argument("--route", choices=["public_notes", "example_essay", "deliverable"], default="deliverable")
    parser.add_argument("--source-map", type=Path)
    parser.add_argument("--allowed", help="Comma-separated extension allow-list.")
    parser.add_argument("--output", type=Path)
    parser.add_argument("--self-test", action="store_true")
    args = parser.parse_args()
    allowed = {item.strip().lower() for item in args.allowed.split(",") if item.strip()} if args.allowed else None
    if args.self_test:
        result = self_test()
    elif args.path:
        result = lint_path(args.path, args.route, args.source_map, allowed)
    else:
        result = {"pass": False, "failures": [{"type": "missing_path_or_self_test"}]}
    text = json.dumps(result, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text + "\n", encoding="utf-8")
    else:
        print(text)
    return 0 if result.get("pass") else 1


if __name__ == "__main__":
    raise SystemExit(main())
